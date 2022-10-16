import datetime
from docx import Document
from docx.shared import Inches

def generate_initial_docx(input_date, input_hour):
    document = Document()

    document.add_heading('Post généré automatiquement par Tansoftware', 0)

    p = document.add_paragraph("🥳 Vous pouvez désormais éditer ce document afin qu'il puisse être posté automatiquement sur LinkedIn.")

    document.add_page_break()
    
    document_name = str(input_date.replace("-", "_")) + "__" + str(input_hour.replace(":", "_"))
    
    print("\n\033[33m-------------------------------------------------------------------------------------------")
    print("\033[33m Le document : " + document_name + ".docx, a été généré dans le répertoire \033[96mscheduled_posts")
    print("\033[33m-------------------------------------------------------------------------------------------")
    print("\n\033[33m Si l'outil ne tourne pas en tâche de fond, pensez à le lancer avec la commande :")
    print("\033[96m python main.py")
    print("\033[33m Puis selectionner le choix : \033[31m2")

    document.save("./scheduled_posts/___" + str(document_name) + ".docx")

def question_date():
    input_date = input("\033[1;32;40m A quelle date voulez vous l'envoyer (au format \033[1;31;40m" + datetime.datetime.now().strftime("%d-%m-%Y") + "\033[1;32;40m) ? : ")
    
    try:
        datetime.datetime.strptime(input_date, '%d-%m-%Y')

        return input_date
    except ValueError:
        print("\033[1;31;40m Le format de la date n'est pas valide, il doit être du type : " + datetime.datetime.now().strftime("%d-%m-%Y\n"))
        return question_date()

def question_hour():
    input_hour = input("\033[1;32;40m A quelle heure voulez vous l'envoyer (au format \033[1;31;40m" + datetime.datetime.now().strftime("%H:%M") + "\033[1;32;40m) ? : ")

    try:
        datetime.datetime.strptime(input_hour, '%H:%M')
        return input_hour
    except ValueError:
        print("\033[1;31;40m Le format de l'heure n'est pas valide, il doit être du type : " + datetime.datetime.now().strftime("%H:%M") + "\n")
        return question_hour()

def questions():
    input_date = question_date()
    input_hour = question_hour()
    
    generate_initial_docx(input_date, input_hour)

def init():
    questions()